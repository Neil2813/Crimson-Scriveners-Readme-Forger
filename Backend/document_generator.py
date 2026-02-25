"""
PDF and DOCX generation from a DocumentModel.
Table colour is applied consistently across both outputs.
"""

from __future__ import annotations
import io
import datetime
import logging
import re

from md_parser import DocumentModel, TableNode, ListNode, CodeBlock
from html_renderer import SAFE_TABLE_COLORS, DEFAULT_COLOR

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Helper — resolve colour palette
# ─────────────────────────────────────────────────────────────────────────────

def _resolve_palette(table_color: str) -> dict:
    """Return the colour palette dict for the given key, falling back to default."""
    return SAFE_TABLE_COLORS.get(table_color, SAFE_TABLE_COLORS[DEFAULT_COLOR])


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Convert '#RRGGBB' or 'RRGGBB' to (r, g, b) integers."""
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(model: DocumentModel, table_color: str = DEFAULT_COLOR) -> bytes:
    """Generate a professional DOCX from DocumentModel using the chosen table colour."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    palette = _resolve_palette(table_color)
    hdr_rgb   = _hex_to_rgb(palette["bg"])
    hdr_text  = _hex_to_rgb(palette["text"])
    hdr_hex   = palette["bg"].lstrip("#").upper()   # for XML shading

    doc = Document()

    # ── Margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # ── Cover ──
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cover_title.add_run(model.title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(17, 24, 39)

    meta = doc.add_paragraph()
    gen_date = datetime.datetime.utcnow().strftime("%B %d, %Y")
    run = meta.add_run(f"Technical Report  •  Generated {gen_date}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()  # spacer

    # ── Sections ──
    heading_sizes  = {1: 16, 2: 13, 3: 11, 4: 10, 5: 9, 6: 9}
    heading_colors = {
        1: (17, 24, 39),
        2: (31, 41, 55),
        3: (55, 65, 81),
        4: (75, 85, 99),
        5: (75, 85, 99),
        6: (75, 85, 99),
    }

    for sec in model.sections:
        if sec.heading:
            lvl = max(1, min(sec.level, 6))
            h = doc.add_paragraph()
            run = h.add_run(sec.heading)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(heading_sizes.get(lvl, 11))
            run.font.color.rgb = RGBColor(*heading_colors.get(lvl, (17, 24, 39)))
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after  = Pt(4)

        for para_text in sec.content:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(30, 30, 30)   # dark text — visible
            p.paragraph_format.space_after = Pt(6)

        for tbl in sec.tables:
            if not tbl.headers and not tbl.rows:
                continue
            total_rows = (1 if tbl.headers else 0) + len(tbl.rows)
            cols = max(len(tbl.headers), max((len(r) for r in tbl.rows), default=0))
            if cols == 0:
                continue
            t = doc.add_table(rows=total_rows, cols=cols)
            t.style = "Table Grid"
            row_idx = 0
            if tbl.headers:
                for ci, h_text in enumerate(tbl.headers):
                    cell = t.rows[0].cells[ci]
                    cell.text = h_text
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(*hdr_text)
                    # User-chosen header background
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), hdr_hex)
                    shading.set(qn("w:color"), "auto")
                    shading.set(qn("w:val"), "clear")
                    cell._tc.get_or_add_tcPr().append(shading)
                row_idx = 1
            for row_data in tbl.rows:
                for ci, cell_text in enumerate(row_data):
                    if ci < cols:
                        cell = t.rows[row_idx].cells[ci]
                        cell.text = cell_text
                        for run in cell.paragraphs[0].runs:
                            run.font.color.rgb = RGBColor(30, 30, 30)
                row_idx += 1
            doc.add_paragraph()  # spacer

        for lst in sec.lists:
            style = "List Number" if lst.ordered else "List Bullet"
            for item in lst.items:
                p = doc.add_paragraph(item, style=style)
                if p.runs:
                    p.runs[0].font.name = "Calibri"
                    p.runs[0].font.size = Pt(10.5)
                    p.runs[0].font.color.rgb = RGBColor(30, 30, 30)

        for cb in sec.code_blocks:
            lang_hint = f"[{cb.language}]  " if cb.language else ""
            p = doc.add_paragraph()
            run = p.add_run(lang_hint + cb.code)
            run.font.name = "Cascadia Code"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(31, 41, 55)
            p.paragraph_format.left_indent = Pt(18)
            pPr = p._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F3F4F6")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:val"), "clear")
            pPr.append(shading)
            doc.add_paragraph()

    # ── Footer ──
    doc.core_properties.description = f"Generated by Crimson Scriveners on {gen_date}"
    doc.core_properties.subject = "Technical Report"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# PDF Generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(html_content: str, table_color: str = DEFAULT_COLOR) -> bytes:
    """Generate PDF from HTML. Uses WeasyPrint when available, falls back to ReportLab."""

    # ── WeasyPrint path (preserves full CSS including the user-chosen colour) ──
    try:
        from weasyprint import HTML as WH
        pdf_bytes = WH(string=html_content).write_pdf()
        return pdf_bytes
    except ImportError:
        pass

    # ── ReportLab fallback: parse HTML and rebuild with chosen palette ─────────
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Preformatted, HRFlowable,
    )
    from reportlab.lib.enums import TA_JUSTIFY
    from bs4 import BeautifulSoup
    import html as html_mod

    palette  = _resolve_palette(table_color)
    hdr_hex  = palette["bg"]       # e.g. "#475569"
    hdr_text = palette["text"]     # e.g. "#ffffff"

    accent_color = colors.HexColor(hdr_hex)
    hdr_txt_color = colors.HexColor(hdr_text)

    dark  = colors.HexColor("#111827")
    gray  = colors.HexColor("#1F2937")   # dark enough on white — fixes visibility
    light = colors.HexColor("#4B5563")   # slightly lighter for meta text, still readable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=cm * 2.54,
        leftMargin=cm * 2.54,
        topMargin=cm * 2.54,
        bottomMargin=cm * 2.54,
    )

    style_title  = ParagraphStyle("DocTitle",  fontSize=24, leading=30, textColor=dark,  fontName="Helvetica-Bold", spaceAfter=4)
    style_meta   = ParagraphStyle("DocMeta",   fontSize=8,  leading=12, textColor=light, fontName="Helvetica",      spaceAfter=16)
    style_h1     = ParagraphStyle("H1",        fontSize=16, leading=22, textColor=dark,  fontName="Helvetica-Bold", spaceBefore=18, spaceAfter=6)
    style_h2     = ParagraphStyle("H2",        fontSize=13, leading=18, textColor=dark,  fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=4)
    style_h3     = ParagraphStyle("H3",        fontSize=11, leading=16, textColor=gray,  fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=3)
    style_body   = ParagraphStyle("Body",      fontSize=10, leading=16, textColor=gray,  fontName="Helvetica",      spaceAfter=8, alignment=TA_JUSTIFY)
    style_bullet = ParagraphStyle("Bullet",    fontSize=10, leading=14, textColor=gray,  fontName="Helvetica",      leftIndent=16, spaceAfter=3, bulletIndent=4)
    style_code   = ParagraphStyle("Code",      fontSize=8,  leading=12, textColor=colors.HexColor("#1F2937"),
                                  fontName="Courier", backColor=colors.HexColor("#F3F4F6"),
                                  leftIndent=12, rightIndent=12, spaceBefore=6, spaceAfter=6)
    style_toc_hdr = ParagraphStyle("TocHdr",  fontSize=9,  leading=14, textColor=accent_color, fontName="Helvetica-Bold", spaceAfter=4)

    soup = BeautifulSoup(html_content, "html.parser")
    story = []

    # Cover
    title_tag = soup.find(class_="doc-title")
    meta_tag  = soup.find(class_="doc-meta")
    if title_tag:
        story.append(Paragraph(html_mod.escape(title_tag.get_text()), style_title))
    if meta_tag:
        story.append(Paragraph(html_mod.escape(meta_tag.get_text()), style_meta))
    story.append(HRFlowable(width="100%", thickness=2.5, color=accent_color, spaceAfter=20))

    # TOC
    toc_div = soup.find(class_="toc")
    if toc_div:
        story.append(Paragraph("Table of Contents", style_toc_hdr))
        for li in toc_div.find_all("li"):
            story.append(Paragraph(f"• {html_mod.escape(li.get_text())}", style_bullet))
        story.append(Spacer(1, 20))

    # Walk sections
    for section_div in soup.find_all(class_="section"):
        for child in section_div.children:
            if not hasattr(child, "name") or not child.name:
                continue
            tag = child.name
            cls = " ".join(child.get("class", []))

            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                lvl = int(tag[1])
                s = {1: style_h1, 2: style_h2, 3: style_h3}.get(lvl, style_h3)
                story.append(Paragraph(html_mod.escape(child.get_text()), s))
                if lvl == 1:
                    story.append(HRFlowable(width="100%", thickness=1.5, color=accent_color, spaceAfter=8))

            elif tag == "p" and "doc-paragraph" in cls:
                story.append(Paragraph(html_mod.escape(child.get_text()), style_body))

            elif tag == "blockquote" or "doc-blockquote" in cls:
                story.append(Paragraph(html_mod.escape(child.get_text()), style_bullet))

            elif "doc-table-wrapper" in cls:
                tbl_el = child.find("table")
                if tbl_el:
                    rows_data = []
                    header_row = tbl_el.find("thead")
                    if header_row:
                        cells = [
                            Paragraph(html_mod.escape(th.get_text()),
                                      ParagraphStyle("TH", fontSize=9, fontName="Helvetica-Bold", textColor=hdr_txt_color))
                            for th in header_row.find_all("th")
                        ]
                        rows_data.append(cells)
                    tbody = tbl_el.find("tbody")
                    if tbody:
                        for tr in tbody.find_all("tr"):
                            cells = [Paragraph(html_mod.escape(td.get_text()), style_body) for td in tr.find_all("td")]
                            if cells:
                                rows_data.append(cells)

                    if rows_data:
                        try:
                            t = Table(rows_data, repeatRows=1 if header_row else 0)
                            ts = [
                                ("BACKGROUND",    (0, 0), (-1, 0), accent_color),
                                ("TEXTCOLOR",     (0, 0), (-1, 0), hdr_txt_color),
                                ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
                                ("FONTSIZE",      (0, 0), (-1, 0), 9),
                                ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.HexColor("#FAFAFA"), colors.HexColor("#F3F4F8")]),
                                ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#E5E7EB")),
                                ("TOPPADDING",    (0, 0), (-1, -1), 7),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                                ("LEFTPADDING",   (0, 0), (-1, -1), 10),
                                ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
                                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                            ]
                            t.setStyle(TableStyle(ts))
                            story.append(t)
                            story.append(Spacer(1, 12))
                        except Exception:
                            pass

            elif "doc-list" in cls:
                for li in child.find_all("li"):
                    story.append(Paragraph(f"• {html_mod.escape(li.get_text())}", style_bullet))
                story.append(Spacer(1, 6))

            elif "doc-code-wrapper" in cls:
                code_el = child.find("code")
                if code_el:
                    lang_el = child.find(class_="doc-code-lang")
                    code_text = code_el.get_text()
                    if lang_el:
                        story.append(Paragraph(
                            f"[{html_mod.escape(lang_el.get_text())}]",
                            ParagraphStyle("CodeLang", fontSize=7, fontName="Helvetica-Bold", textColor=gray),
                        ))
                    try:
                        story.append(Preformatted(code_text, style_code))
                    except Exception:
                        story.append(Paragraph(html_mod.escape(code_text[:500]), style_body))
                    story.append(Spacer(1, 8))

    story.append(Spacer(1, 40))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E5E7EB")))
    story.append(Paragraph("Generated by Crimson Scriveners Readme Forger", style_meta))

    doc.build(story)
    buf.seek(0)
    return buf.read()
